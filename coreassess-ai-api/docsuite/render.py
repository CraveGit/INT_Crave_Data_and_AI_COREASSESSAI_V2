import os
import re
import html
import markdown
from bs4 import BeautifulSoup
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import RichText

from .style import (ASSETS, BODY_FONT, HEADING_FONT, MONO_FONT, HEADING_COLOR,
                    BODY_COLOR, MUTED_COLOR, BORDER_COLOR, BODY_SIZE, TABLE_SIZE,
                    LINE_SPACING, PARA_SPACE_AFTER, HEADING_SIZES,
                    HEADING_SPACE_BEFORE, HEADING_SPACE_AFTER, applyBaseStyles)

_BULLET_GLYPHS = ('', 'o', '')   # Symbol dot, Courier o, Wingdings square
_BULLET_RE = re.compile(r'^\s*[-*+]\s+(.*)$')
_ORDERED_RE = re.compile(r'^\s*\d+[.)]\s+(.*)$')


def make_bold(text):
    rt = RichText()
    rt.add(str(text), bold=True, italic=True)
    return rt


def _is_heading(line):
    return re.match(r'^#{1,6}\s', line) is not None


def _is_table_row(line):
    return re.match(r'^\|.*\|$', line) is not None


def _is_separator_row(cells):
    return all(re.fullmatch(r':?-{2,}:?', c.strip()) for c in cells if c.strip())


def _bullet_depth(line):
    return min((len(line) - len(line.lstrip(' '))) // 2, 3)


# Inline markdown -> runs: **bold**, *italic*, `code`.
def _process_text(paragraph, text, color=BODY_COLOR, size=BODY_SIZE, font=BODY_FONT):
    text = html.unescape(text)
    for seg in re.split(r'(\*\*.+?\*\*|`[^`]+`|(?<!\*)\*[^*]+\*(?!\*))', text):
        if not seg:
            continue
        run = paragraph.add_run()
        if seg.startswith('**') and seg.endswith('**'):
            run.text, run.bold = seg[2:-2], True
            run.font.name = font
        elif seg.startswith('`') and seg.endswith('`'):
            run.text = seg[1:-1]
            run.font.name = MONO_FONT
        elif seg.startswith('*') and seg.endswith('*'):
            run.text, run.italic = seg[1:-1], True
            run.font.name = font
        else:
            run.text = seg
            run.font.name = font
        run.font.color.rgb = color
        run.font.size = Pt(size)
    return paragraph


# Split only comma lists of bare identifiers (MARA, MARC); never prose.
def _is_value_list(txt):
    if ',' not in txt or len(txt) > 200 or '\n' in txt:
        return False
    parts = [p.strip() for p in txt.split(',')]
    if len(parts) < 2 or any(not p for p in parts):
        return False
    if re.search(r'\b(and|or|the|to|with|for|of|is|are|which|that)\b', txt, re.I):
        return False
    if any(p.endswith('.') for p in parts):
        return False
    return all(len(p.split()) == 1 and re.fullmatch(r'[A-Za-z0-9_./-]+', p) for p in parts)


def _cell_borders(cell):
    tc_pr = cell._element.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), BORDER_COLOR)
        borders.append(el)
    tc_pr.append(borders)


def _fill_cell(cell, text, header=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _cell_borders(cell)
    cell.text = ''
    for i, line in enumerate(text.split('\n') if '\n' in text else [text]):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        fmt = para.paragraph_format
        fmt.space_before = Pt(2)
        fmt.space_after = Pt(2)
        fmt.line_spacing = 1.15
        if line.startswith('• '):
            fmt.left_indent = Inches(0.14)
            line = line[2:]
            _process_text(para, '• ', BODY_COLOR, TABLE_SIZE)
        for run in _process_text(para, line, BODY_COLOR, TABLE_SIZE).runs:
            if header:
                run.bold = True


def _add_table(doc, rows):
    if not rows:
        return
    header = rows[0]
    data = [r for r in rows[1:] if not _is_separator_row(r)]
    width = max(len(header), max((len(r) for r in data), default=0))
    table = doc.add_table(rows=len(data) + 1, cols=width)
    table.style = 'Table Grid'
    table.autofit = True

    for j in range(width):
        # No shading: the reference export separates the header with weight, not fill.
        _fill_cell(table.cell(0, j), header[j].strip() if j < len(header) else '', header=True)
    for i, row in enumerate(data):
        for j in range(width):
            txt = row[j].strip() if j < len(row) else ''
            txt = '\n'.join(l.replace('- ', '• ') for l in txt.split('<br>'))
            if _is_value_list(txt):
                txt = '\n'.join(f'• {v.strip()}' for v in txt.split(','))
            _fill_cell(table.cell(i + 1, j), txt)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _numbering_part(doc):
    try:
        return doc.part.numbering_part
    except (KeyError, NotImplementedError, AttributeError):
        return None


# The abstract numbering (the FORMAT: decimal vs bullet, indents, glyphs) is built
# once per doc and cached. Concrete <w:num> instances reference it.
def _abstract_num(doc, ordered):
    part = _numbering_part(doc)
    if part is None:
        return None, None
    numbering = part.element
    key = '_abs_ordered' if ordered else '_abs_bullet'
    if getattr(doc, key, None) is not None:
        return numbering, getattr(doc, key)

    ids = [int(a.get(qn('w:abstractNumId'))) for a in numbering.findall(qn('w:abstractNum'))]
    abstract_id = max(ids) + 1 if ids else 0
    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abstract_id))
    for level in range(4):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level))
        start = OxmlElement('w:start'); start.set(qn('w:val'), '1'); lvl.append(start)
        fmt = OxmlElement('w:numFmt')
        fmt.set(qn('w:val'), 'decimal' if ordered else 'bullet'); lvl.append(fmt)
        text = OxmlElement('w:lvlText')
        text.set(qn('w:val'), f'%{level + 1}.' if ordered else _BULLET_GLYPHS[level % 3])
        lvl.append(text)
        jc = OxmlElement('w:lvlJc'); jc.set(qn('w:val'), 'left'); lvl.append(jc)
        ppr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(360 + 340 * level))
        ind.set(qn('w:hanging'), '360')
        ppr.append(ind); lvl.append(ppr)
        if not ordered:
            rpr = OxmlElement('w:rPr')
            fonts = OxmlElement('w:rFonts')
            glyph_font = 'Symbol' if level % 3 == 0 else ('Courier New' if level % 3 == 1 else 'Wingdings')
            fonts.set(qn('w:ascii'), glyph_font); fonts.set(qn('w:hAnsi'), glyph_font)
            rpr.append(fonts); lvl.append(rpr)
        abstract.append(lvl)
    numbering.insert(0, abstract)
    setattr(doc, key, abstract_id)
    return numbering, abstract_id


# A FRESH <w:num> for each list. Two lists with different numIds keep independent
# counters, so every ordered list RESTARTS at 1 instead of continuing the previous
# list's count across the whole document (the "1,2,3 ... then 22,23,24" bug).
def _new_list_num(doc, ordered):
    numbering, abstract_id = _abstract_num(doc, ordered)
    if numbering is None:
        return None
    num_ids = [int(n.get(qn('w:numId'))) for n in numbering.findall(qn('w:num'))]
    num_id = max(num_ids) + 1 if num_ids else 1
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    ref = OxmlElement('w:abstractNumId'); ref.set(qn('w:val'), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def _add_bullet(doc, text, depth, ordered, num_id):
    para = doc.add_paragraph()
    if num_id is not None:
        ppr = para._p.get_or_add_pPr()
        num_pr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), str(depth)); num_pr.append(ilvl)
        nid = OxmlElement('w:numId'); nid.set(qn('w:val'), str(num_id)); num_pr.append(nid)
        ppr.append(num_pr)
        _process_text(para, text)
    else:
        _process_text(para, ('• ' if not ordered else '- ') + text)
    fmt = para.paragraph_format
    fmt.left_indent = Inches(0.28 + 0.24 * depth)
    fmt.space_after = Pt(3)
    fmt.line_spacing = LINE_SPACING
    return para


def _style_heading(doc, text, level):
    heading = doc.add_heading('', level=level)
    fmt = heading.paragraph_format
    fmt.space_before = Pt(HEADING_SPACE_BEFORE.get(level, 8))
    fmt.space_after = Pt(HEADING_SPACE_AFTER.get(level, 5))
    fmt.keep_with_next = True
    if level == 1:
        fmt.page_break_before = True   # every chapter opens a page
    _process_text(heading, text, HEADING_COLOR, HEADING_SIZES.get(level, 11), HEADING_FONT)
    for run in heading.runs:
        run.bold = True
        run.font.name = HEADING_FONT
    return heading


def _add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.name = BODY_FONT
    run.font.color.rgb = MUTED_COLOR
    para.paragraph_format.space_after = Pt(10)


def _strip_code_fence(lines):
    """Drop the wrapper fence the model puts around the whole answer, but keep
    inner fences: pseudo-code blocks are content, not packaging."""
    fences = [i for i, l in enumerate(lines) if re.match(r'^\s*```', l)]
    if len(fences) >= 2 and fences[0] <= 1 and fences[-1] >= len(lines) - 2:
        drop = {fences[0], fences[-1]}
        return [l for i, l in enumerate(lines) if i not in drop]
    return list(lines)


def _add_code_block(doc, code_lines):
    for line in code_lines:
        para = doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.left_indent = Inches(0.22)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.1
        run = para.add_run(line)
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = BODY_COLOR
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _html_to_markdown(text):
    from markdownify import markdownify
    return markdownify(text, heading_style="ATX")


def _looks_like_html(lines):
    joined = "\n".join(lines)
    tags = re.findall(r'</?(?:div|html|body|table|tr|td|th|h[1-6]|p|span|head|meta|style|ul|li)\b',
                      joined, re.I)
    return len(tags) >= 15


def markdown_into_doc(lines, doc, diagrams=None):
    """Render markdown into an open docx. `diagrams` maps heading text -> image
    stream inserted directly under that heading."""
    lines = _strip_code_fence(list(lines))
    if _looks_like_html(lines):
        lines = _html_to_markdown("\n".join(lines)).splitlines()
    applyBaseStyles(doc)
    diagrams = diagrams or {}
    placed_figures = set()
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            _add_table(doc, table_rows)
            table_rows = []

    code_buffer, in_code = [], False
    # Current list's numId per kind. A run of same-kind items shares one numId; any
    # interruption (other kind, heading, paragraph, table, code) resets it to None so
    # the NEXT list allocates a fresh numId and restarts at 1.
    ordered_num = bullet_num = None

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith('```'):
            ordered_num = bullet_num = None
            if in_code:
                _add_code_block(doc, code_buffer)
                code_buffer, in_code = [], False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue

        if _is_table_row(stripped):
            ordered_num = bullet_num = None
            table_rows.append([html.unescape(c.strip()) for c in stripped.strip('|').split('|')])
            continue
        flush_table()

        if not stripped:
            continue
        if re.fullmatch(r'(\*\s*){3,}|(-\s*){3,}|(_\s*){3,}', stripped):
            continue

        if _is_heading(stripped):
            ordered_num = bullet_num = None
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            _style_heading(doc, text, min(level, 6))
            # Only chapter openers carry figures, and each figure is placed once:
            # subsections often repeat the chapter title.
            key = re.sub(r'^[\d.\s]+', '', text).strip().lower()
            if level == 1 and key in diagrams and key not in placed_figures:
                stream = diagrams[key]
                if stream is not None:
                    stream.seek(0)
                    doc.add_picture(stream, width=Inches(6.2))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Computed outside the f-string: Python 3.11 (the CF runtime)
                    # rejects a backslash inside an f-string expression.
                    caption = re.sub(r'^[\d.\s]+', '', text).strip()
                    _add_caption(doc, f"Figure: {caption}")
                    placed_figures.add(key)
            continue

        m = _BULLET_RE.match(line)
        if m:
            ordered_num = None                       # a bullet ends any ordered run
            if bullet_num is None:
                bullet_num = _new_list_num(doc, False)
            _add_bullet(doc, m.group(1).strip(), _bullet_depth(line), False, bullet_num)
            continue
        m = _ORDERED_RE.match(line)
        if m:
            bullet_num = None                        # an ordered item ends any bullet run
            if ordered_num is None:                  # new ordered list -> fresh numId, restarts at 1
                ordered_num = _new_list_num(doc, True)
            _add_bullet(doc, m.group(1).strip(), _bullet_depth(line), True, ordered_num)
            continue

        ordered_num = bullet_num = None
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(PARA_SPACE_AFTER)
        para.paragraph_format.line_spacing = LINE_SPACING
        _process_text(para, stripped)

    if in_code and code_buffer:
        _add_code_block(doc, code_buffer)
    flush_table()


def markdown_to_styled_html(md_text, font=None, size="11px"):
    body = markdown.markdown(md_text, extensions=['tables'])
    css = (
        "<style>"
        "@import url('https://fonts.googleapis.com/css2?"
        "family=Instrument+Sans:wght@400..700&display=swap');"
        "@import url('https://api.fontshare.com/v2/css?f[]=general-sans@400,500,600,700&display=swap');"
        "body{font-family:'Instrument Sans',sans-serif;font-size:%s;line-height:1.4;"
        "letter-spacing:.02rem;margin:40px;color:#111;max-width:1000px}"
        "h1,h2,h3,h4,h5,h6{font-family:'General Sans','Instrument Sans',sans-serif;font-weight:700;color:#494558;"
        "margin:1.2em 0 .4em}"
        "h1{font-size:1.9em}h2{font-size:1.55em}h3{font-size:1.3em}h4{font-size:1.15em}"
        "p{margin:.9rem 0}"
        "table{border-collapse:collapse;width:100%%;margin:1rem 0;font-size:.9em}"
        "th,td{border:1px solid #d8d8d4;padding:8px;text-align:left;vertical-align:top}"
        "th{font-weight:700}"
        "code{font-family:Consolas,monospace;font-size:.9em}"
        "ul,ol{margin:.6rem 0 .6rem 1.2rem}li{margin:.2rem 0}"
        "hr{border:0;border-top:1px solid #aaa}"
        "</style>"
    ) % size
    full = f"<html><head><meta charset='UTF-8'>{css}</head><body>{body}</body></html>"
    # Do NOT prettify: prettify() inserts newlines/indentation INSIDE inline elements
    # (<strong>, <code>, ...). Those newlines survive the round-trip to .docx
    # (markdownify -> markdown_into_doc), where each inline fragment then lands on its
    # own paragraph -- code spans and bold text broke onto separate lines throughout
    # the document. decode() keeps compact HTML: each paragraph on one logical line.
    return BeautifulSoup(full, 'html.parser').decode()
